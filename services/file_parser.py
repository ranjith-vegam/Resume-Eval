import re
import io
from pathlib import Path
from config import MAX_RESUME_CHARS


def _extract_pdf(data: bytes) -> str:
    import fitz
    doc = fitz.open(stream=data, filetype="pdf")
    pages = [page.get_text("text") for page in doc]
    return "\n".join(pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells if c.text.strip()))
    return "\n".join(parts)


def _extract_txt(data: bytes) -> str:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        text = _extract_pdf(data)
    elif suffix in (".docx", ".doc"):
        text = _extract_docx(data)
    else:
        text = _extract_txt(data)
    return text[:MAX_RESUME_CHARS]


_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}")
_NAME_PATTERNS = [
    re.compile(r"^([A-Z][a-z]+(?: [A-Z][a-z]+){1,3})"),
    re.compile(r"Name\s*[:\-]\s*([A-Z][a-z]+(?: [A-Z][a-z]+){1,3})", re.IGNORECASE),
]
_PHONE_RE = re.compile(r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{3,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{0,4}")


def _find_phone(text: str) -> str | None:
    for m in _PHONE_RE.finditer(text[:2000]):
        digits = re.sub(r"\D", "", m.group(0))
        if 7 <= len(digits) <= 15:
            return m.group(0).strip()
    return None


def extract_candidate_info(text: str) -> tuple[str | None, str | None, str | None]:
    email_match = _EMAIL_RE.search(text[:2000])
    email = email_match.group(0) if email_match else None

    name = None
    for pat in _NAME_PATTERNS:
        m = pat.search(text[:500])
        if m:
            name = m.group(1)
            break

    phone = _find_phone(text)

    return name, email, phone
